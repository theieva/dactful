"""
Core correctness tests - the fixtures that break naive redaction tools.
If these pass, the engine is trustworthy. Everything else is plumbing.
"""

import os
import sys

import pytest
from docx import Document
from docx.oxml.ns import qn

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from app.matching import term_rule, redact_text, find_matches  # noqa: E402
from app.docx_redact import redact_docx, extract_text  # noqa: E402
from app.restore import restore_docx, restore_text  # noqa: E402


# --- entity set reused across tests ----------------------------------------
def entities():
    return [
        ("Acme Corp Inc", "[[CLIENT1_NAME]]"),
        ("Acme Credit", "[[CLIENT1_DBA]]"),
        ("Clienty Corp", "[[CLIENT2_NAME]]"),
        ("Clienty", "[[CLIENT2_DBA]]"),
        ("Banky Bank", "[[SPONSOR_BANK1]]"),
    ]


def rules():
    return [term_rule(t, tag) for t, tag in entities()]


# --- text-level matching ----------------------------------------------------
def test_substring_trap():
    # "Clienty Corp" must win over "Clienty" - never "[[CLIENT2_DBA]] Corp".
    out = redact_text("Clienty Corp signed with Clienty today.", rules())
    assert out == "[[CLIENT2_NAME]] signed with [[CLIENT2_DBA]] today."


def test_overlapping_entities_longest_first():
    out = redact_text(
        "Acme Corp Inc, doing business as Acme Credit, is not Acme alone.",
        rules() + [term_rule("Acme", "[[CLIENT1_SHORT]]")],
    )
    assert "[[CLIENT1_NAME]]" in out
    assert "[[CLIENT1_DBA]]" in out
    assert "[[CLIENT1_SHORT]] alone" in out
    assert "Corp Inc" not in out  # longer match was not partially eaten


def test_possessive_preserved():
    out = redact_text("That is Clienty's headquarters.", rules())
    assert out == "That is [[CLIENT2_DBA]]'s headquarters."


def test_case_insensitive():
    out = redact_text("BANKY BANK and banky bank both.", rules())
    assert out == "[[SPONSOR_BANK1]] and [[SPONSOR_BANK1]] both."


def test_line_broken_name():
    # Name split across a line break still matches (flexible whitespace).
    out = redact_text("Signed by Acme\nCorp Inc yesterday.", rules())
    assert out == "Signed by [[CLIENT1_NAME]] yesterday."


def test_no_false_plural_match():
    # "Acme" should NOT fire inside "Acmes" (no accidental over-redaction).
    out = redact_text("The Acmes are unrelated.", [term_rule("Acme", "[[X]]")])
    assert out == "The Acmes are unrelated."


# --- docx run fragmentation -------------------------------------------------
def _para_with_runs(doc, run_texts):
    p = doc.add_paragraph()
    for rt in run_texts:
        p.add_run(rt)
    return p


def test_docx_run_fragmentation(tmp_path):
    doc = Document()
    # "Acme Corp Inc" deliberately split across runs, incl. a mid-word split.
    _para_with_runs(doc, ["Client is ", "Ac", "me Corp", " Inc", " today."])
    _para_with_runs(doc, ["Also ", "Clienty", " Corp here."])
    src = tmp_path / "in.docx"
    doc.save(str(src))

    out = tmp_path / "out.docx"
    redact_docx(str(src), str(out), rules())
    text = extract_text(str(out))

    assert "[[CLIENT1_NAME]]" in text
    assert "[[CLIENT2_NAME]]" in text
    assert "Acme" not in text  # nothing leaked despite the split runs
    assert "Clienty" not in text


def test_docx_hidden_surfaces(tmp_path):
    doc = Document()
    doc.add_paragraph("Body mentions Acme Corp Inc.")
    section = doc.sections[0]
    section.header.paragraphs[0].add_run("Header: Banky Bank confidential")
    section.footer.paragraphs[0].add_run("Footer prepared for Clienty Corp")
    doc.core_properties.author = "Acme Corp Inc"
    doc.core_properties.title = "Deal with Clienty Corp"
    src = tmp_path / "in.docx"
    doc.save(str(src))

    out = tmp_path / "out.docx"
    redact_docx(str(src), str(out), rules())
    text = extract_text(str(out))

    for leaked in ("Acme", "Banky Bank", "Clienty"):
        assert leaked not in text, f"leaked from a hidden surface: {leaked}"
    assert "[[SPONSOR_BANK1]]" in text  # header redacted
    assert "[[CLIENT1_NAME]]" in text   # core-props author redacted


def test_docx_preserves_untouched_formatting(tmp_path):
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Bold intro ").bold = True
    p.add_run("Acme Corp Inc")
    p.add_run(" plain tail")
    src = tmp_path / "in.docx"
    doc.save(str(src))
    out = tmp_path / "out.docx"
    redact_docx(str(src), str(out), rules())

    rd = Document(str(out))
    runs = rd.paragraphs[0].runs
    assert runs[0].bold is True
    assert "Bold intro" in runs[0].text
    assert "[[CLIENT1_NAME]]" in "".join(r.text for r in runs)


# --- round trip -------------------------------------------------------------
def test_round_trip_text():
    original = "Acme Corp Inc and Clienty Corp met Banky Bank about Clienty."
    ents = entities()
    red = redact_text(original, [term_rule(t, tag) for t, tag in ents])
    # The mapping only carries terms that were actually used (as the app builds it).
    mapping = [{"tag": tag, "value": t} for t, tag in ents if t in original]
    restored, report = restore_text(red, mapping)
    assert restored == original
    assert report.leftover == []


def test_round_trip_docx(tmp_path):
    doc = Document()
    doc.add_paragraph("Acme Corp Inc partnered with Clienty Corp.")
    doc.add_paragraph("Banky Bank is the sponsor for Clienty.")
    src = tmp_path / "in.docx"
    doc.save(str(src))

    red = tmp_path / "red.docx"
    redact_docx(str(src), str(red), rules())
    doc_text = "Acme Corp Inc partnered with Clienty Corp. Banky Bank is the sponsor for Clienty."
    mapping = [{"tag": tag, "value": t} for t, tag in entities() if t in doc_text]
    fin = tmp_path / "fin.docx"
    report = restore_docx(str(red), str(fin), mapping)

    restored_text = extract_text(str(fin))
    assert "Acme Corp Inc" in restored_text
    assert "Clienty Corp" in restored_text
    assert "Banky Bank" in restored_text
    assert report.leftover == []


def test_restore_repairs_mangled_tags():
    mapping = [{"tag": "[[CLIENT1_NAME]]", "value": "Acme Corp Inc"}]
    text = "Draft for **[[CLIENT1_NAME]]** and [[ CLIENT1_NAME ]] and [[Client1 Name]]."
    restored, report = restore_text(text, mapping)
    assert "Acme Corp Inc" in restored
    assert "[[" not in restored  # all three variants repaired
    assert report.mangled["[[CLIENT1_NAME]]"] == 3


def test_restore_reports_leftover_unknown_tag():
    # A tag with no value (invented/altered by the AI) is reported as leftover.
    mapping = [{"tag": "[[CLIENT1_NAME]]", "value": "Acme Corp Inc"}]
    restored, report = restore_text("Hi [[CLIENT1_NAME]] and [[UNKNOWN_9]].", mapping)
    assert "Acme Corp Inc" in restored
    assert report.total == 1
    assert report.leftover == ["[[UNKNOWN_9]]"]


if __name__ == "__main__":
    sys.exit(pytest.main([__file__, "-v"]))
