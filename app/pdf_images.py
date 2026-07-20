"""
Extract and index the images embedded in a PDF, so the user can choose to keep
or drop each one.

Each image is normalized to PNG, given a small thumbnail (base64 data URI for
the review UI), and OCR-scanned on-device so we can warn when a picture looks
like it holds text (a name, an account number). Dactful redacts TEXT; it cannot
black out the inside of an image, so keeping one sends it to the AI as-is. The
warning helps the user review before deciding.
"""

from __future__ import annotations

import base64
import io
import os
from typing import Callable, Dict, List, Optional

# Skip decorative junk (icons, spacers, hairlines).
_MIN_DIM = 40
_THUMB = (260, 260)


def available() -> bool:
    try:
        import fitz  # noqa: F401
        from PIL import Image  # noqa: F401

        return True
    except Exception:
        return False


def _classify(png_path: str) -> str:
    """'sensitive' | 'text' | 'none' based on on-device OCR of the image."""
    from .detect import detect_patterns
    from .ocr import available as ocr_available, image_to_text

    if not ocr_available():
        return "unknown"
    try:
        text = image_to_text(png_path)
    except Exception:
        return "none"
    if not text.strip():
        return "none"
    return "sensitive" if detect_patterns(text) else "text"


_MAX_IMAGES = 24


def _cluster_rects(rects, gap):
    """Merge rectangles that overlap or sit within `gap` of each other."""
    import fitz

    boxes = [fitz.Rect(r) for r in rects if r.width >= 1 and r.height >= 1]
    changed = True
    while changed:
        changed = False
        i = 0
        while i < len(boxes):
            j = i + 1
            while j < len(boxes):
                a = boxes[i]
                a_exp = fitz.Rect(a.x0 - gap, a.y0 - gap, a.x1 + gap, a.y1 + gap)
                if a_exp.intersects(boxes[j]):
                    boxes[i] = a | boxes[j]
                    boxes.pop(j)
                    changed = True
                else:
                    j += 1
            i += 1
    return boxes


def _vector_regions(page):
    """Cluster the page's vector drawings (charts, logos) into renderable regions.
    These are drawn shapes, not embedded images, so get_images() misses them."""
    import fitz

    rects = []
    for d in page.get_drawings():
        r = d.get("rect")
        if r is None:
            continue
        r = fitz.Rect(r)
        if r.width < 8 or r.height < 8:  # skip hairlines / borders
            continue
        rects.append(r)
    if not rects:
        return []

    page_area = abs(page.rect.width * page.rect.height)
    out = []
    # gap is generous so a chart's bars/segments group into one image rather
    # than fragmenting into many.
    for c in _cluster_rects(rects, gap=26):
        if c.width < 50 or c.height < 40:                       # too small
            continue
        if abs(c.width * c.height) > 0.85 * page_area:          # near full page
            continue
        if c.width > 0.92 * page.rect.width and c.height < 26:  # divider band
            continue
        out.append(c)
    return out


def extract_and_index(pdf_path: str, out_dir: str) -> List[Dict]:
    """Return [{id, path, thumb, warn, page, width, height}] for each embedded
    image AND each vector-graphic region (charts, logos), or [] if unavailable."""
    if not available():
        return []
    import fitz
    from PIL import Image

    os.makedirs(out_dir, exist_ok=True)
    try:
        doc = fitz.open(pdf_path)
    except Exception:
        return []

    out: List[Dict] = []

    def add(im, page_num):
        w, h = im.size
        if w < _MIN_DIM or h < _MIN_DIM:
            return
        idx = len(out)
        png_path = os.path.join(out_dir, f"img_{idx}.png")
        if im.mode not in ("RGB", "RGBA"):
            im = im.convert("RGB")
        im.save(png_path, "PNG")
        thumb = im.copy()
        thumb.thumbnail(_THUMB)
        buf = io.BytesIO()
        thumb.convert("RGB").save(buf, "PNG")
        out.append({
            "id": idx,
            "path": png_path,
            "thumb": "data:image/png;base64," + base64.b64encode(buf.getvalue()).decode(),
            "warn": _classify(png_path),
            "page": page_num + 1,
            "width": w,
            "height": h,
        })

    seen = set()
    try:
        for page_num in range(len(doc)):
            if len(out) >= _MAX_IMAGES:
                break
            page = doc[page_num]
            # 1. embedded raster images
            for img in page.get_images(full=True):
                xref = img[0]
                if xref in seen:
                    continue
                seen.add(xref)
                try:
                    im = Image.open(io.BytesIO(doc.extract_image(xref)["image"]))
                except Exception:
                    continue
                add(im, page_num)
            # 2. vector-graphic regions, rendered to PNG
            for rect in _vector_regions(page):
                if len(out) >= _MAX_IMAGES:
                    break
                try:
                    pix = page.get_pixmap(clip=rect, dpi=150)
                    im = Image.open(io.BytesIO(pix.tobytes("png")))
                except Exception:
                    continue
                add(im, page_num)
    finally:
        doc.close()
    return out


def embed_images(docx_path: str, image_paths: List[str]) -> None:
    """Append the kept images to the end of the redacted .docx."""
    if not image_paths:
        return
    from docx import Document
    from docx.shared import Inches

    doc = Document(docx_path)
    doc.add_page_break()
    doc.add_heading("Images from the original document", level=2)
    for p in image_paths:
        try:
            doc.add_picture(p, width=Inches(6))
        except Exception:
            pass
    doc.save(docx_path)
